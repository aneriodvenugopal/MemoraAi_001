"""
MemoraAI Iteration 13 backend tests

Covers:
- Tenant admin login (phone 8888888888)
- Staff Members API: GET, POST (with generated default password), DELETE
- Logs API: GET /api/memoraai/logs returns events[]/counts
- Knowledge extract: PDF, DOCX, unsupported binary -> 422
- Simplified CRM: POST /api/memoraai/crm/leads and /contacts with minimal payload
- Automation stat data: /rules + /corrections/summary
"""
import io
import os
import uuid

import pytest
import requests

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://retomerp-memora.preview.emergentagent.com").rstrip("/")
API = f"{BASE_URL}/api"
TENANT_PHONE = "8888888888"
TENANT_PASSWORD = "admin123"


@pytest.fixture(scope="module")
def token():
    r = requests.post(f"{API}/auth/login", json={"phone": TENANT_PHONE, "password": TENANT_PASSWORD}, timeout=30)
    assert r.status_code == 200, f"login failed {r.status_code} {r.text}"
    data = r.json()
    t = data.get("access_token") or data.get("token")
    assert t, f"no token in login response: {data}"
    return t


@pytest.fixture(scope="module")
def auth_headers(token):
    return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}


@pytest.fixture(scope="module")
def auth_headers_form(token):
    return {"Authorization": f"Bearer {token}"}


# ---------- Login ----------
def test_tenant_login_returns_token_and_role():
    r = requests.post(f"{API}/auth/login", json={"phone": TENANT_PHONE, "password": TENANT_PASSWORD}, timeout=30)
    assert r.status_code == 200
    j = r.json()
    assert j.get("access_token") or j.get("token")
    user = j.get("user") or {}
    assert user.get("role") == "tenant_admin"


# ---------- Staff ----------
class TestStaff:
    created_id = None

    def test_list_staff(self, auth_headers):
        r = requests.get(f"{API}/memoraai/staff", headers=auth_headers, timeout=30)
        assert r.status_code == 200, r.text
        j = r.json()
        assert "staff" in j and isinstance(j["staff"], list)
        assert j["total"] >= 1  # at least tenant admin
        # Ensure tenant admin present
        roles = [s.get("role") for s in j["staff"]]
        assert "tenant_admin" in roles

    def test_add_staff_and_verify(self, auth_headers):
        unique_phone = f"9{uuid.uuid4().int % 10**9:09d}"  # random 10-digit phone
        payload = {
            "name": "TEST_Staff_Iter13",
            "phone": unique_phone,
            "permissions": ["content", "leads"],
        }
        r = requests.post(f"{API}/memoraai/staff", headers=auth_headers, json=payload, timeout=30)
        assert r.status_code == 200, r.text
        j = r.json()
        assert j.get("success") is True
        assert j.get("default_password") == "memora@123"
        assert j.get("staff_id")
        TestStaff.created_id = j["staff_id"]

        # Verify via GET
        r2 = requests.get(f"{API}/memoraai/staff", headers=auth_headers, timeout=30)
        assert r2.status_code == 200
        staff = r2.json()["staff"]
        phones = [s.get("phone") for s in staff]
        assert unique_phone in phones

    def test_add_staff_duplicate_phone(self, auth_headers):
        payload = {"name": "TEST_Dup", "phone": TENANT_PHONE}
        r = requests.post(f"{API}/memoraai/staff", headers=auth_headers, json=payload, timeout=30)
        assert r.status_code == 400

    def test_delete_staff(self, auth_headers):
        assert TestStaff.created_id, "create test should have run first"
        r = requests.delete(f"{API}/memoraai/staff/{TestStaff.created_id}", headers=auth_headers, timeout=30)
        assert r.status_code == 200, r.text
        assert r.json().get("success") is True


# ---------- Logs ----------
def test_logs_returns_events_shape(auth_headers):
    r = requests.get(f"{API}/memoraai/logs", headers=auth_headers, timeout=30)
    assert r.status_code == 200, r.text
    j = r.json()
    assert "events" in j and isinstance(j["events"], list)
    assert "counts" in j and isinstance(j["counts"], dict)
    assert "generated_at" in j
    # Each event should have the unified shape
    for e in j["events"][:5]:
        assert "kind" in e
        assert "title" in e
        assert "timestamp" in e


def test_logs_filter_by_kind(auth_headers):
    for k in ["ai", "correction", "content", "lead", "staff"]:
        r = requests.get(f"{API}/memoraai/logs", params={"kind": k, "limit": 20}, headers=auth_headers, timeout=30)
        assert r.status_code == 200, f"{k}: {r.text}"
        kinds = {e["kind"] for e in r.json()["events"]}
        assert kinds.issubset({k}) or not kinds  # all events must be that kind (or empty)


# ---------- Knowledge extract ----------
def _minimal_pdf_bytes():
    # Tiny valid PDF with text "Hello"
    return (
        b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 144]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\nBT /F1 24 Tf 50 100 Td (Hello World KB) Tj ET\nendstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000060 00000 n \n"
        b"0000000110 00000 n \n0000000210 00000 n \n0000000290 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n360\n%%EOF"
    )


def _docx_bytes():
    from docx import Document
    d = Document()
    d.add_paragraph("TEST_DOCX content for MemoraAI iteration 13 knowledge base.")
    d.add_paragraph("Second paragraph with pricing info.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_knowledge_extract_pdf(auth_headers_form):
    files = {"file": ("hello.pdf", _minimal_pdf_bytes(), "application/pdf")}
    r = requests.post(f"{API}/memoraai/knowledge/extract", headers=auth_headers_form, files=files, timeout=60)
    assert r.status_code == 200, r.text
    j = r.json()
    assert j.get("success") is True
    assert j.get("char_count", 0) > 0
    assert "Hello World" in j.get("extracted_text", "") or len(j["extracted_text"]) > 0


def test_knowledge_extract_docx(auth_headers_form):
    files = {
        "file": (
            "sample.docx",
            _docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    r = requests.post(f"{API}/memoraai/knowledge/extract", headers=auth_headers_form, files=files, timeout=60)
    assert r.status_code == 200, r.text
    j = r.json()
    assert j.get("success") is True
    assert "TEST_DOCX content" in j.get("extracted_text", "")


def test_knowledge_extract_unsupported_binary_returns_422(auth_headers_form):
    # Random binary bytes with no readable text - should 422
    binary = bytes([i % 256 for i in range(4096)])
    files = {"file": ("garbage.bin", binary, "application/octet-stream")}
    r = requests.post(f"{API}/memoraai/knowledge/extract", headers=auth_headers_form, files=files, timeout=60)
    assert r.status_code == 422, f"expected 422 got {r.status_code}: {r.text}"
    assert "Could not read" in r.json().get("detail", "")


# ---------- CRM simplified ----------
def test_create_lead_simplified(auth_headers):
    payload = {
        "name": "TEST_Lead_Iter13",
        "phone": f"7{uuid.uuid4().int % 10**9:09d}",
        "source": "walk_in",
        "notes": "Met them at a friend's party, interested in pricing.",
    }
    r = requests.post(f"{API}/memoraai/crm/leads", headers=auth_headers, json=payload, timeout=30)
    assert r.status_code in (200, 201), r.text
    j = r.json()
    # Try various common shapes
    lead = j.get("lead") or j
    assert (lead.get("name") == payload["name"]) or j.get("success")


def test_create_contact_simplified(auth_headers):
    payload = {
        "name": "TEST_Contact_Iter13",
        "phone": f"6{uuid.uuid4().int % 10**9:09d}",
        "notes": "Existing customer, friendly.",
    }
    r = requests.post(f"{API}/memoraai/crm/contacts", headers=auth_headers, json=payload, timeout=30)
    assert r.status_code in (200, 201), r.text


# ---------- Automation stat sources ----------
def test_rules_endpoint(auth_headers):
    r = requests.get(f"{API}/memoraai/rules", headers=auth_headers, timeout=30)
    assert r.status_code == 200, r.text


    def test_corrections_summary(self, auth_headers):
        # Endpoint is /stats/summary (the /summary alone hits /{correction_id})
        r = requests.get(f"{API}/memoraai/corrections/stats/summary", headers=auth_headers, timeout=30)
        assert r.status_code == 200, r.text
        j = r.json()
        assert isinstance(j, dict)


def test_corrections_summary(auth_headers):
    r = requests.get(f"{API}/memoraai/corrections/stats/summary", headers=auth_headers, timeout=30)
    assert r.status_code == 200, r.text
